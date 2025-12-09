import streamlit as st
import pandas as pd
import requests
import time
import json
import re
import io

# --- IMPORT SAFETY CHECK ---
try:
    import google.generativeai as genai
    HAS_GENAI = True
except ImportError:
    HAS_GENAI = False

# FPDF removed as requested

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from odf.opendocument import OpenDocumentText
    from odf.style import Style, TextProperties, ParagraphProperties
    from odf.text import H, P, Span
    HAS_ODT = True
except ImportError:
    HAS_ODT = False

# --- API KEY SETUP ---
HARDCODED_KEY = "" 

# --- CONFIGURATION ---
st.set_page_config(
    page_title="LeetCode Study Helper",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- FALLBACK DATASET ---
FALLBACK_DB = [
    {"id": 1, "title": "Two Sum", "titleSlug": "two-sum", "difficulty": "Easy", "url": "https://leetcode.com/problems/two-sum/", "desc": "Given an array of integers nums and an integer target, return indices of the two numbers such that they add up to target."},
    {"id": 2, "title": "Add Two Numbers", "titleSlug": "add-two-numbers", "difficulty": "Medium", "url": "https://leetcode.com/problems/add-two-numbers/", "desc": "Add two numbers represented by linked lists."},
    {"id": 3, "title": "Longest Substring Without Repeating Characters", "titleSlug": "longest-substring-without-repeating-characters", "difficulty": "Medium", "url": "https://leetcode.com/problems/longest-substring-without-repeating-characters/", "desc": "Find length of longest substring with unique chars."}
]

# --- LIVE FETCHING FUNCTIONS ---
@st.cache_data(ttl=3600)
def fetch_problem_list(skip=0, limit=50):
    url = "https://leetcode.com/graphql"
    query = """
    query problemsetQuestionList($categorySlug: String, $limit: Int, $skip: Int, $filters: QuestionListFilterInput) {
      problemsetQuestionList: questionList(
        categorySlug: $categorySlug
        limit: $limit
        skip: $skip
        filters: $filters
      ) {
        questions: data {
          frontendQuestionId: questionFrontendId
          title
          titleSlug
          difficulty
        }
      }
    }
    """
    variables = {"categorySlug": "all-code-essentials", "limit": limit, "skip": skip, "filters": {}}
    try:
        response = requests.post(url, json={'query': query, 'variables': variables}, timeout=5)
        if response.status_code == 200:
            data = response.json()
            questions = data['data']['problemsetQuestionList']['questions']
            formatted = []
            for q in questions:
                formatted.append({
                    "id": q['frontendQuestionId'],
                    "title": q['title'],
                    "titleSlug": q['titleSlug'],
                    "difficulty": q['difficulty'],
                    "url": f"https://leetcode.com/problems/{q['titleSlug']}/",
                    "desc": "Fetching description..." 
                })
            return formatted
    except Exception as e:
        print(f"Fetch Error: {e}")
    return None

@st.cache_data(ttl=3600)
def fetch_problem_content(title_slug):
    url = "https://leetcode.com/graphql"
    query = """
    query questionContent($titleSlug: String!) {
      question(titleSlug: $titleSlug) {
        content
      }
    }
    """
    variables = {"titleSlug": title_slug}
    try:
        response = requests.post(url, json={'query': query, 'variables': variables}, timeout=5)
        if response.status_code == 200:
            data = response.json()
            return data['data']['question']['content']
    except:
        return None
    return None

# --- AI HELPERS ---
def generate_gemini_text(api_key, prompt):
    if not HAS_GENAI: return "Error: Library missing"
    genai.configure(api_key=api_key)
    # Updated model list to prioritize Gemini 2.5 Flash
    models_to_try = [
        'gemini-2.5-flash', 
        'gemini-2.5-flash-preview-09-2025', 
        'gemini-1.5-flash'
    ]
    
    error_log = []
    for model_name in models_to_try:
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            error_log.append(f"{model_name}: {str(e)}")
            continue
            
    # Return detailed errors for debugging
    return f"Error: AI generation failed.\nDebug Info: {json.dumps(error_log, indent=2)}"

# --- CONTENT GENERATOR (Separated from File Creation) ---
def generate_study_content(problems_queue, api_key):
    generated_data = []
    progress_bar = st.progress(0)
    status_text = st.empty()

    for i, prob in enumerate(problems_queue):
        status_text.text(f"Analyzing {prob['title']} ({i+1}/{len(problems_queue)})...")
        
        # 1. Fetch Details if missing
        desc_content = prob['desc']
        if desc_content == "Fetching description...":
            fetched = fetch_problem_content(prob['titleSlug'])
            if fetched: desc_content = fetched
        
        # 2. AI Prompt
        prompt = f"""
        Generate a study report for the LeetCode problem "{prob['title']}" (ID: {prob['id']}).
        Problem Description: {desc_content[:2000]} (truncated)
        
        Output format:
        [UNDERSTANDING]
        Write a paragraph explaining the problem and the logic to solve it efficiently (Time Complexity). Start with "My understanding: ...". DO NOT use pseudo-code here, use natural language.
        
        [EXAMPLES]
        Provide 2-3 input/output examples. Format: "Input: ... Output: ..."
        
        [SOLUTION]
        Provide the complete C++ solution code class.
        IMPORTANT: Provide ONLY the C++ code. Do not add comments (`//`) explaining the logic inside the code block. Just the clean, working solution.
        """
        
        ai_response = generate_gemini_text(api_key, prompt)
        
        # Parse Response
        understanding = ""
        examples = ""
        solution = ""
        
        if ai_response.startswith("Error:"):
            understanding = ai_response
        else:
            parts = ai_response.split('[EXAMPLES]')
            if len(parts) > 0:
                understanding = parts[0].replace('[UNDERSTANDING]', '').strip()
            
            if len(parts) > 1:
                sub_parts = parts[1].split('[SOLUTION]')
                examples = sub_parts[0].strip()
                if len(sub_parts) > 1:
                    solution = sub_parts[1].strip().replace('```cpp', '').replace('```', '')

        generated_data.append({
            "id": prob['id'],
            "title": prob['title'],
            "url": prob['url'],
            "difficulty": prob['difficulty'],
            "understanding": understanding,
            "examples": examples,
            "solution": solution
        })
        
        progress_bar.progress((i + 1) / len(problems_queue))
    
    status_text.text("Content Generated! Ready to export.")
    time.sleep(1)
    status_text.empty()
    progress_bar.empty()
    
    return generated_data

# --- FILE EXPORTERS ---

# 1. DOCX
def create_docx_bytes(report_data):
    if not HAS_DOCX: return None
    doc = Document()
    doc.add_heading('Submission of Self-Practice Evidence', 0)
    
    for i, data in enumerate(report_data):
        # Main Title (The Problem Name)
        doc.add_heading(f"{i+1}. {data['title']}", level=1)
        
        # 1. Problem {id}
        doc.add_paragraph(f"1. Problem {data['id']}")
        
        # 2. URL: {url}
        doc.add_paragraph(f"2. URL: {data['url']}")
        
        # 3. Level of problem: {difficulty}
        doc.add_paragraph(f"3. Level of problem: {data['difficulty']}")
        
        # 4. My understanding:
        p_und = doc.add_paragraph()
        run = p_und.add_run("4. My understanding:")
        run.bold = True
        p_und.add_run(f" {data['understanding']}")
        
        # Examples listed under understanding
        doc.add_paragraph(data['examples'])
        
        # 5. Solution:
        p_sol = doc.add_paragraph()
        run = p_sol.add_run("5. Solution:")
        run.bold = True

        p_code = doc.add_paragraph(data['solution'])
        for run in p_code.runs:
            run.font.name = 'Courier New'
        if not p_code.runs:
            run = p_code.add_run(data['solution'])
            run.font.name = 'Courier New'

        # 6. Screenshot:
        p_shot = doc.add_paragraph()
        run = p_shot.add_run("6. Screenshot:")
        run.bold = True
        
        p_placeholder = doc.add_paragraph("ADD YOUR SCREENSHOT HERE")
        p_placeholder.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in p_placeholder.runs:
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            
        doc.add_page_break()
        
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# 2. ODT
def create_odt_bytes(report_data):
    if not HAS_ODT: return None
    textdoc = OpenDocumentText()
    
    # Styles
    h1style = Style(name="Heading 1", family="paragraph")
    h1style.addElement(TextProperties(attributes={'fontsize':"14pt",'fontweight':"bold"}))
    textdoc.styles.addElement(h1style)

    boldstyle = Style(name="BoldLabel", family="text")
    boldstyle.addElement(TextProperties(attributes={'fontweight':"bold"}))
    textdoc.styles.addElement(boldstyle)

    codestyle = Style(name="Code", family="paragraph")
    codestyle.addElement(TextProperties(attributes={'fontfamily':"Courier New", 'fontsize':"10pt"}))
    textdoc.styles.addElement(codestyle)

    # Updated: Separated ParagraphProperties (for alignment) and TextProperties (for color/style)
    placeholderstyle = Style(name="Placeholder", family="paragraph")
    placeholderstyle.addElement(ParagraphProperties(attributes={'textalign':"center"}))
    placeholderstyle.addElement(TextProperties(attributes={'fontstyle':"italic", 'color':"#888888"}))
    textdoc.styles.addElement(placeholderstyle)
    
    # Title
    textdoc.text.addElement(H(outlinelevel=1, text="Submission of Self-Practice Evidence"))
    
    for i, data in enumerate(report_data):
        # Main Title
        textdoc.text.addElement(H(outlinelevel=2, text=f"{i+1}. {data['title']}"))
        
        # 1. Problem {id}
        textdoc.text.addElement(P(text=f"1. Problem {data['id']}"))
        
        # 2. URL: {url}
        textdoc.text.addElement(P(text=f"2. URL: {data['url']}"))
        
        # 3. Level of problem: {difficulty}
        textdoc.text.addElement(P(text=f"3. Level of problem: {data['difficulty']}"))
        
        # 4. My understanding:
        p_und = P()
        span_und = Span(stylename=boldstyle, text="4. My understanding:")
        p_und.addElement(span_und)
        p_und.addText(f" {data['understanding']}")
        textdoc.text.addElement(p_und)
        
        # Examples
        textdoc.text.addElement(P(text=data['examples']))
        
        # 5. Solution:
        p_sol_label = P()
        span_sol = Span(stylename=boldstyle, text="5. Solution:")
        p_sol_label.addElement(span_sol)
        textdoc.text.addElement(p_sol_label)

        p_code = P(stylename=codestyle, text=str(data['solution']))
        textdoc.text.addElement(p_code)

        # 6. Screenshot:
        p_shot_label = P()
        span_shot = Span(stylename=boldstyle, text="6. Screenshot:")
        p_shot_label.addElement(span_shot)
        textdoc.text.addElement(p_shot_label)

        textdoc.text.addElement(P(stylename=placeholderstyle, text="ADD YOUR SCREENSHOT HERE"))
        
    buffer = io.BytesIO()
    textdoc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# --- INITIALIZATION ---
if 'page_skip' not in st.session_state: st.session_state['page_skip'] = 0
if 'pdf_queue' not in st.session_state: st.session_state['pdf_queue'] = []
if 'generated_report_data' not in st.session_state: st.session_state['generated_report_data'] = None
if 'current_problem_idx' not in st.session_state: st.session_state['current_problem_idx'] = 0
if 'hints' not in st.session_state: st.session_state['hints'] = {}

# --- SIDEBAR ---
with st.sidebar:
    st.header("⚙️ Configuration")
    if HAS_GENAI:
        if HARDCODED_KEY:
            gemini_api_key = HARDCODED_KEY
            st.success("✅ Hardcoded Key Active")
        else:
            gemini_api_key = st.text_input("Gemini API Key", type="password")
    else:
        st.warning("Google GenAI Lib Missing")
        gemini_api_key = None

    st.markdown("---")
    
    # --- REPORT QUEUE MANAGEMENT ---
    st.header("📝 Study Report Queue")
    if len(st.session_state['pdf_queue']) == 0:
        st.info("Queue is empty. Add problems from the main view.")
    else:
        st.success(f"**{len(st.session_state['pdf_queue'])}** problems in queue.")
        
        with st.expander("View Queue"):
            for p in st.session_state['pdf_queue']:
                st.write(f"- {p['title']}")
            if st.button("Clear Queue"):
                st.session_state['pdf_queue'] = []
                st.session_state['generated_report_data'] = None
                st.rerun()

        st.markdown("### 1. Generate Content")
        if st.button("🧠 Analyze & Prepare Report"):
            if not gemini_api_key:
                st.error("Need API Key")
            else:
                with st.spinner("Asking AI to analyze problems..."):
                    data = generate_study_content(st.session_state['pdf_queue'], gemini_api_key)
                    st.session_state['generated_report_data'] = data
                    st.rerun()

        if st.session_state['generated_report_data']:
            st.markdown("### 2. Export As")
            
            # DOCX
            if HAS_DOCX:
                docx_bytes = create_docx_bytes(st.session_state['generated_report_data'])
                st.download_button("📝 Download Word (.docx)", docx_bytes, "Study_Report.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            else:
                st.error("Install `python-docx`")

            # ODT
            if HAS_ODT:
                odt_bytes = create_odt_bytes(st.session_state['generated_report_data'])
                st.download_button("📓 Download ODT", odt_bytes, "Study_Report.odt", "application/vnd.oasis.opendocument.text")
            else:
                st.error("Install `odfpy`")

    st.markdown("---")
    if st.button("🔄 Reload App"):
        st.cache_data.clear()
        st.rerun()

# --- MAIN LOGIC ---
PAGE_SIZE = 50
with st.spinner("Loading Problems..."):
    live_data = fetch_problem_list(skip=st.session_state['page_skip'], limit=PAGE_SIZE)

current_db = live_data if live_data else FALLBACK_DB

# 2. LIST VIEW
st.title("🧩 LeetCode AI Study Helper")

col_prev, col_info, col_next = st.columns([1, 4, 1])
with col_prev:
    if st.button("⬅️ Prev"):
        if st.session_state['page_skip'] >= PAGE_SIZE:
            st.session_state['page_skip'] -= PAGE_SIZE
            st.session_state['current_problem_idx'] = 0
            st.rerun()
with col_next:
    if st.button("Next ➡️"):
        st.session_state['page_skip'] += PAGE_SIZE
        st.session_state['current_problem_idx'] = 0
        st.rerun()

options = [f"{p['id']}: {p['title']}" for p in current_db]
safe_index = min(st.session_state['current_problem_idx'], len(options) - 1)
selected_option = st.selectbox("Select Problem:", options, index=safe_index)

if selected_option in options:
    new_idx = options.index(selected_option)
    if st.session_state['current_problem_idx'] != new_idx:
        st.session_state['current_problem_idx'] = new_idx
        st.rerun()

# --- ACTIVE PROBLEM WORKSPACE ---
if len(current_db) > 0:
    active_problem = current_db[st.session_state['current_problem_idx']]
    p_id = active_problem['id']

    if p_id not in st.session_state['hints']:
        st.session_state['hints'][p_id] = []

    # HEADER & ACTIONS
    c1, c2 = st.columns([3, 1])
    with c1:
        st.header(f"{active_problem['title']}")
        st.caption(f"**Difficulty:** {active_problem['difficulty']}")
    with c2:
        in_queue = any(p['id'] == active_problem['id'] for p in st.session_state['pdf_queue'])
        
        if in_queue:
            st.button("✅ In Queue", disabled=True)
        else:
            if st.button("➕ Add to Report Queue"):
                st.session_state['pdf_queue'].append(active_problem)
                st.session_state['generated_report_data'] = None 
                st.rerun()

    # FETCH DESCRIPTION
    if active_problem['desc'] == "Fetching description...":
        with st.spinner("Fetching details..."):
            content = fetch_problem_content(active_problem['titleSlug'])
            if content: active_problem['desc'] = content
    
    # TABS FOR VIEW
    tab1, tab2, tab3 = st.tabs(["🖥️ Workspace", "💡 Hints", "🔓 AI Solution"])
    
    with tab1:
        st.info(f"Link: {active_problem['url']}")
        try:
            st.components.v1.iframe(active_problem['url'], height=700, scrolling=True)
        except:
            st.error("Embed blocked. Use external link.")

    with tab2:
        st.subheader("💡 Progressive Hints")
        
        current_hints = st.session_state['hints'][p_id]
        if not current_hints:
            st.info("No hints generated yet. Click the button below to start.")
        
        for idx, h in enumerate(current_hints):
            st.info(f"**Hint {idx+1}:** {h}")
            
        if gemini_api_key:
            if st.button("🔍 Get Next Hint"):
                hint_prompt = f"""
                Provide a short, progressive hint for the LeetCode problem "{active_problem['title']}".
                The user is stuck.
                
                Previous hints given: {current_hints}
                
                Goal: Nudge them towards the efficient solution (Time Complexity) without writing the full code or giving the answer away immediately.
                If no hints given yet, give a conceptual starting point.
                """
                with st.spinner("Thinking of a clue..."):
                    new_hint = generate_gemini_text(gemini_api_key, hint_prompt)
                    st.session_state['hints'][p_id].append(new_hint)
                    st.rerun()
        else:
            st.warning("Enter API Key to get hints.")

    with tab3:
        if gemini_api_key:
            if st.button("✨ Generate C++ Solution"):
                prompt = f"Solve LeetCode '{active_problem['title']}' in C++. Just code."
                with st.spinner("Coding..."):
                    code = generate_gemini_text(gemini_api_key, prompt)
                    st.code(code.replace('```cpp','').replace('```',''), language='cpp')
        else:
            st.warning("Enter API Key to generate solutions.")
